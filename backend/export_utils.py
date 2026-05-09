import io
import os

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ------------------------------------------------------------------ Word helpers

def _set_rtl(para):
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    _set_rtl(h)
    return h


def _para(doc, text, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.add_run(text)
    _set_rtl(p)
    return p


# ------------------------------------------------------------------ Word export

def generate_word(blueprint: dict) -> bytes:
    doc = Document()

    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)

    title = doc.add_heading('מסמך אפיון פונקציונלי', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Spec document sections
    for section in blueprint.get('spec_document', {}).get('sections', []):
        _heading(doc, section.get('title', ''), 1)
        for line in section.get('content', '').split('\n'):
            s = line.strip()
            if not s:
                continue
            if s.startswith('- ') or s.startswith('• '):
                _para(doc, s[2:], 'List Bullet')
            elif s.startswith('**') and s.endswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(s.strip('*'))
                run.bold = True
                _set_rtl(p)
            else:
                _para(doc, s.lstrip('#').strip())

    # Entities
    entities = blueprint.get('entities', [])
    if entities:
        _heading(doc, 'ישויות המערכת', 1)
        for entity in entities:
            _heading(doc, entity.get('name', ''), 2)
            fields = entity.get('fields', [])
            if fields:
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                for i, label in enumerate(['שם שדה', 'סוג', 'חובה', 'תיאור']):
                    hdr[i].text = label
                    if hdr[i].paragraphs[0].runs:
                        hdr[i].paragraphs[0].runs[0].bold = True
                    hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for f in fields:
                    row = table.add_row().cells
                    row[0].text = f.get('name', '')
                    row[1].text = f.get('type', '')
                    row[2].text = 'כן' if f.get('required') else 'לא'
                    row[3].text = f.get('description', '')
                    for cell in row:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                doc.add_paragraph()

    # Workflows
    workflows = blueprint.get('workflows', [])
    if workflows:
        _heading(doc, 'תהליכים עסקיים', 1)
        for wf in workflows:
            _heading(doc, wf.get('name', ''), 2)
            if wf.get('steps'):
                _para(doc, 'שלבים:')
                for step in wf['steps']:
                    _para(doc, step, 'List Number')
            if wf.get('constraints'):
                _para(doc, 'אילוצים:')
                for c in wf['constraints']:
                    _para(doc, c, 'List Bullet')
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ------------------------------------------------------------------ PDF export

def generate_pdf(blueprint: dict) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from bidi.algorithm import get_display

    # Register a Hebrew-capable font (Arial on Windows, fallback to Helvetica)
    font_name = 'Helvetica'
    for fp in ['C:/Windows/Fonts/arial.ttf', 'C:/Windows/Fonts/Arial.ttf']:
        if os.path.exists(fp):
            try:
                pdfmetrics.registerFont(TTFont('ArialHeb', fp))
                font_name = 'ArialHeb'
            except Exception:
                pass
            break

    def rtl(text: str) -> str:
        return get_display(str(text)) if text else ''

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle('H1', parent=styles['Heading1'], fontName=font_name, alignment=2, fontSize=16)
    h2 = ParagraphStyle('H2', parent=styles['Heading2'], fontName=font_name, alignment=2, fontSize=13)
    norm = ParagraphStyle('Norm', parent=styles['Normal'], fontName=font_name, alignment=2, fontSize=10, leading=15)

    buf = io.BytesIO()
    doc_pdf = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=2 * cm, leftMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
    )
    story = []

    story.append(Paragraph(rtl('מסמך אפיון פונקציונלי'), h1))
    story.append(Spacer(1, 0.5 * cm))

    for section in blueprint.get('spec_document', {}).get('sections', []):
        story.append(Paragraph(rtl(section.get('title', '')), h1))
        for line in section.get('content', '').split('\n'):
            s = line.strip().lstrip('#-• ').strip()
            if s:
                story.append(Paragraph(rtl(s), norm))
        story.append(Spacer(1, 0.3 * cm))

    if blueprint.get('entities'):
        story.append(Paragraph(rtl('ישויות המערכת'), h1))
        for entity in blueprint['entities']:
            story.append(Paragraph(rtl(entity.get('name', '')), h2))
            fields = entity.get('fields', [])
            if fields:
                data = [[rtl('שם שדה'), rtl('סוג'), rtl('חובה'), rtl('תיאור')]]
                for f in fields:
                    data.append([
                        rtl(f.get('name', '')),
                        rtl(f.get('type', '')),
                        rtl('כן' if f.get('required') else 'לא'),
                        rtl(f.get('description', '')),
                    ])
                t = Table(data, colWidths=[3 * cm, 2.5 * cm, 1.8 * cm, 9 * cm])
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#dbeafe')),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
                    ('FONTNAME', (0, 0), (-1, -1), font_name),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8fafc')]),
                ]))
                story.append(t)
            story.append(Spacer(1, 0.3 * cm))

    if blueprint.get('workflows'):
        story.append(Paragraph(rtl('תהליכים עסקיים'), h1))
        for wf in blueprint['workflows']:
            story.append(Paragraph(rtl(wf.get('name', '')), h2))
            for step in wf.get('steps', []):
                story.append(Paragraph(f'• {rtl(step)}', norm))
            for c in wf.get('constraints', []):
                story.append(Paragraph(f'- {rtl(c)}', norm))
            story.append(Spacer(1, 0.2 * cm))

    doc_pdf.build(story)
    buf.seek(0)
    return buf.read()
